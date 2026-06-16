"""
ATS NLP Engine
==============
Keyword extraction, TF-IDF vectorization, cosine similarity scoring,
and detailed gap analysis — all without requiring a pretrained spaCy model.

Uses:
  - spaCy blank tokenizer  (fast, no model download needed)
  - NLTK stopwords + PorterStemmer
  - scikit-learn TfidfVectorizer + cosine_similarity
  - python-docx for .docx extraction
"""

import re
import io
import math
import string
from collections import Counter
from typing import Optional

import spacy
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from docx import Document


# ── Setup ─────────────────────────────────────────────────────────────────────
nlp      = spacy.blank("en")          # fast tokenizer, no model needed
stemmer  = PorterStemmer()
STOPWORDS = set(stopwords.words("english")) | {
    # Domain-specific noise words that aren't useful as skills
    "experience", "year", "years", "work", "working", "role", "team", "ability",
    "skill", "skills", "knowledge", "strong", "good", "excellent", "required",
    "preferred", "plus", "etc", "including", "also", "must", "will", "well",
    "looking", "candidate", "responsibilities", "qualifications", "job", "position",
    "company", "opportunity", "environment", "using", "use", "used", "based",
    "help", "develop", "ensure", "provide", "build", "create", "make", "need",
    "different", "various", "key", "new", "high", "large", "understand", "related",
    # JD structural words that aren't skills
    "expertise", "requirements", "requirement", "frameworks", "framework",
    "proficiency", "familiarity", "databases", "database", "apis", "api",
    "pipelines", "pipeline", "actions", "solving", "senior", "junior", "mid",
    "preferred", "bonus", "nice", "minimum", "least", "plus", "years",
    "responsibilities", "qualifications", "required", "preferred", "benefits",
    "description", "overview", "about", "apply", "email", "resume", "application",
    "candidate", "hire", "hiring", "position", "role", "full", "time", "part",
    "remote", "onsite", "hybrid", "office", "location", "salary", "compensation",
    "competitive", "benefits", "health", "dental", "vision", "equity", "stock",
    "startup", "team", "member", "join", "growing", "fast", "paced",
}

# Tech skills that are acronyms / short — must not be filtered by length
KNOWN_TECH_TERMS = {
    "ai", "ml", "nlp", "api", "sql", "aws", "gcp", "ci", "cd", "ui", "ux",
    "ios", "etl", "bi", "qa", "oop", "sdk", "cli", "crm", "erp", "rpa",
    "css", "html", "php", "git", "r", "go", "c", "c++", "c#", "vb",
}

# Multi-word tech phrases to detect (bigram whitelist patterns)
TECH_PHRASES = [
    r"machine learning", r"deep learning", r"natural language processing",
    r"computer vision", r"data science", r"data engineering", r"data analysis",
    r"software engineering", r"software development", r"web development",
    r"cloud computing", r"big data", r"neural network", r"large language model",
    r"generative ai", r"feature engineering", r"model deployment",
    r"agile methodology", r"test driven development", r"continuous integration",
    r"continuous deployment", r"micro.?services?", r"object.?oriented",
    r"version control", r"open source", r"real.?time", r"end.?to.?end",
    r"restful? api", r"graph.?ql", r"type.?script", r"next\.?js", r"node\.?js",
    r"react\.?js", r"vue\.?js", r"spring boot", r"fast.?api", r"django rest",
    r"scikit.?learn", r"tensor.?flow", r"py.?torch", r"hugging.?face",
    r"apache kafka", r"apache spark", r"apache airflow", r"elastic.?search",
    r"power bi", r"look.?er", r"table.?au",
]
_PHRASE_RE = re.compile(
    r"(?:" + "|".join(TECH_PHRASES) + r")",
    re.IGNORECASE,
)


# ── Text extraction ────────────────────────────────────────────────────────────
def extract_text_from_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def clean_text(text: str) -> str:
    """Normalize whitespace, remove control chars, keep punctuation for phrase detection."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)   # strip non-ASCII
    text = re.sub(r"[ \t]+", " ", text)             # collapse spaces
    return text.strip()


# ── Keyword extraction ────────────────────────────────────────────────────────
def _stem(word: str) -> str:
    return stemmer.stem(word.lower())


def extract_keywords(text: str, top_n: int = 60) -> list[str]:
    """
    Extract meaningful keywords from text using:
      1. Multi-word tech phrase detection (regex)
      2. spaCy blank tokenizer for single-word tokens
      3. PorterStemmer for normalization
      4. TF-IDF-style frequency weighting
    Returns a de-duplicated list of canonical keyword strings (not stems).
    """
    text = clean_text(text)
    text_lower = text.lower()

    found_phrases = []
    phrase_spans  = []
    for m in _PHRASE_RE.finditer(text_lower):
        phrase = re.sub(r"\s+", " ", m.group()).strip()
        found_phrases.append(phrase)
        phrase_spans.append((m.start(), m.end()))

    # Mask matched phrase regions so we don't re-tokenize them
    masked = list(text_lower)
    for s, e in phrase_spans:
        for i in range(s, e):
            masked[i] = " "
    masked_text = "".join(masked)

    # Single-word token extraction via spaCy blank tokenizer
    doc = nlp(masked_text)
    single_words = []
    for token in doc:
        w = token.text.strip().lower()
        if not w:
            continue
        # Keep known tech terms even if short
        if w in KNOWN_TECH_TERMS:
            single_words.append(w)
            continue
        # Filter: alpha or alpha+symbols (C++, .NET etc.), min length 3
        if not re.match(r"^[a-z][a-z\-\+\#\.]{2,}$", w):
            continue
        if w in STOPWORDS:
            continue
        if w in string.punctuation:
            continue
        single_words.append(w)

    # Combine and count frequency
    all_terms = found_phrases + single_words
    freq = Counter(all_terms)

    # Build canonical term map: term → stem (for dedup)
    seen_stems: dict[str, str] = {}  # stem → best canonical form
    for term, count in freq.most_common(top_n * 3):
        stem_key = " ".join(_stem(w) for w in term.split()) if " " in term else _stem(term)
        if stem_key not in seen_stems:
            seen_stems[stem_key] = term
        else:
            # Prefer the longer / more specific form
            if len(term) > len(seen_stems[stem_key]):
                seen_stems[stem_key] = term

    # Sort by frequency descending
    ranked = sorted(seen_stems.values(), key=lambda t: freq[t], reverse=True)
    return ranked[:top_n]


# ── Scoring ───────────────────────────────────────────────────────────────────
def compute_tfidf_similarity(jd_text: str, resume_text: str) -> float:
    """TF-IDF cosine similarity between job description and resume."""
    vectorizer = TfidfVectorizer(
        ngram_range=(1, 2),
        stop_words="english",
        max_features=5000,
        sublinear_tf=True,
    )
    try:
        tfidf = vectorizer.fit_transform([jd_text, resume_text])
        sim   = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
        return float(sim)
    except Exception:
        return 0.0


def keyword_match_score(jd_keywords: list[str], resume_text: str) -> dict:
    """
    Check which JD keywords appear in the resume (using stem matching).
    Returns matched, missing, and match percentage.
    """
    resume_lower = resume_text.lower()
    resume_stems = {_stem(w) for w in re.findall(r"[a-zA-Z][a-zA-Z\-\+\#\.]{1,}", resume_lower)}

    matched, missing = [], []
    for kw in jd_keywords:
        kw_stem = " ".join(_stem(w) for w in kw.split())
        # Check literal substring first (catches "FastAPI", "PyTorch" etc.)
        literal_match = kw.lower() in resume_lower
        # Then stem-based fallback for single words
        stem_match = (kw_stem in resume_stems) if " " not in kw else (kw_stem in resume_lower)
        if literal_match or stem_match:
            matched.append(kw)
        else:
            missing.append(kw)

    pct = len(matched) / len(jd_keywords) * 100 if jd_keywords else 0
    return {"matched": matched, "missing": missing, "match_pct": round(pct, 1)}


def categorize_keywords(keywords: list[str]) -> dict[str, list[str]]:
    """Heuristic categorization of keywords into groups."""
    PROGRAMMING = {
        "python", "java", "javascript", "typescript", "golang", "rust", "c",
        "c++", "c#", "ruby", "scala", "kotlin", "swift", "php", "r", "matlab",
        "perl", "bash", "shell", "powershell", "sql", "nosql",
    }
    FRAMEWORKS = {
        "react", "angular", "vue", "django", "flask", "fastapi", "spring",
        "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
        "express", "rails", "laravel", "next.js", "node.js", "react.js",
        "vue.js", "spring boot", "fastapi", "hugging face",
    }
    CLOUD_INFRA = {
        "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ansible",
        "jenkins", "gitlab", "github", "ci/cd", "linux", "nginx", "redis",
        "kafka", "elasticsearch", "airflow", "spark", "hadoop",
    }
    DATA = {
        "machine learning", "deep learning", "nlp", "natural language processing",
        "computer vision", "data science", "data engineering", "data analysis",
        "sql", "nosql", "postgresql", "mysql", "mongodb", "snowflake", "bigquery",
        "tableau", "power bi", "looker", "spark", "airflow", "etl", "data pipeline",
    }
    SOFT = {
        "leadership", "communication", "collaboration", "agile", "scrum",
        "problem solving", "critical thinking", "project management", "mentoring",
    }

    cats: dict[str, list[str]] = {
        "Programming Languages": [],
        "Frameworks & Libraries": [],
        "Cloud & Infrastructure": [],
        "Data & ML": [],
        "Soft Skills": [],
        "Other Technical": [],
    }
    for kw in keywords:
        kw_l = kw.lower()
        if any(kw_l == p or kw_l.startswith(p) for p in PROGRAMMING):
            cats["Programming Languages"].append(kw)
        elif any(kw_l == f or f in kw_l for f in FRAMEWORKS):
            cats["Frameworks & Libraries"].append(kw)
        elif any(kw_l == c or c in kw_l for c in CLOUD_INFRA):
            cats["Cloud & Infrastructure"].append(kw)
        elif any(kw_l == d or d in kw_l for d in DATA):
            cats["Data & ML"].append(kw)
        elif any(kw_l == s or s in kw_l for s in SOFT):
            cats["Soft Skills"].append(kw)
        else:
            cats["Other Technical"].append(kw)
    return {k: v for k, v in cats.items() if v}


def compute_ats_score(
    jd_text: str,
    resume_text: str,
    jd_keywords: list[str],
) -> dict:
    """
    Composite ATS score:
      40% — TF-IDF cosine similarity
      45% — keyword match percentage
      15% — structural bonus (contact info, sections detected)
    """
    cosine_sim   = compute_tfidf_similarity(jd_text, resume_text)
    kw_result    = keyword_match_score(jd_keywords, resume_text)

    # Structural bonus
    resume_lower = resume_text.lower()
    struct_score = 0.0
    indicators = {
        "email":       bool(re.search(r"[\w.]+@[\w.]+\.\w+", resume_text)),
        "phone":       bool(re.search(r"\+?\d[\d\s\-\(\)]{8,}", resume_text)),
        "linkedin":    "linkedin" in resume_lower,
        "education":   any(w in resume_lower for w in ["education", "university", "college", "degree", "bachelor", "master", "phd"]),
        "experience":  any(w in resume_lower for w in ["experience", "work history", "employment", "worked at"]),
        "projects":    any(w in resume_lower for w in ["project", "built", "developed", "created", "designed"]),
        "skills_section": any(w in resume_lower for w in ["skills", "technologies", "competencies", "expertise"]),
        "achievements": any(w in resume_lower for w in ["achieved", "improved", "increased", "reduced", "delivered", "launched", "%"]),
    }
    struct_score = sum(indicators.values()) / len(indicators)

    # Weighted composite
    ats_score = (
        0.40 * cosine_sim * 100 +
        0.45 * kw_result["match_pct"] +
        0.15 * struct_score * 100
    )
    ats_score = min(round(ats_score, 1), 100.0)

    # Grade
    if ats_score >= 80:
        grade, verdict = "A", "Excellent match — strong ATS pass likelihood"
    elif ats_score >= 65:
        grade, verdict = "B", "Good match — minor gaps to address"
    elif ats_score >= 50:
        grade, verdict = "C", "Moderate match — several improvements needed"
    elif ats_score >= 35:
        grade, verdict = "D", "Weak match — significant skill gaps"
    else:
        grade, verdict = "F", "Poor match — resume needs substantial revision"

    # Actionable tips
    tips = []
    if kw_result["match_pct"] < 60:
        tips.append("Add more keywords from the job description directly into your resume.")
    if not indicators["achievements"]:
        tips.append("Quantify achievements with numbers (%, $, x faster) — ATS and humans love metrics.")
    if not indicators["skills_section"]:
        tips.append("Add a dedicated 'Skills' section for fast ATS parsing.")
    if not indicators["linkedin"]:
        tips.append("Include your LinkedIn URL — many ATS systems look for it.")
    if cosine_sim < 0.15:
        tips.append("Mirror the job description's language more closely in your resume.")
    if len(kw_result["missing"]) > 10:
        tips.append(f"Incorporate at least {min(5, len(kw_result['missing']))} of the missing keywords.")

    return {
        "ats_score":      ats_score,
        "grade":          grade,
        "verdict":        verdict,
        "cosine_sim":     round(cosine_sim * 100, 1),
        "keyword_match":  kw_result["match_pct"],
        "struct_score":   round(struct_score * 100, 1),
        "matched_keywords": kw_result["matched"],
        "missing_keywords": kw_result["missing"],
        "jd_keywords":    jd_keywords,
        "struct_indicators": indicators,
        "tips":           tips,
        "categories":     categorize_keywords(kw_result["missing"]),
    }


# ── Main scoring entry point ──────────────────────────────────────────────────
def score_resume(
    jd_text: str,
    resume_text: str,
    jd_file: Optional[bytes] = None,
    resume_file: Optional[bytes] = None,
    jd_filename: str = "",
    resume_filename: str = "",
) -> dict:
    """
    Full scoring pipeline. Accepts plain text or .docx bytes.
    """
    # Parse uploads if provided
    if jd_file and jd_filename.lower().endswith(".docx"):
        jd_text = extract_text_from_docx(jd_file)
    if resume_file and resume_filename.lower().endswith(".docx"):
        resume_text = extract_text_from_docx(resume_file)

    jd_text     = clean_text(jd_text or "")
    resume_text = clean_text(resume_text or "")

    if len(jd_text) < 30:
        raise ValueError("Job description is too short. Please provide more detail.")
    if len(resume_text) < 30:
        raise ValueError("Resume is too short. Please provide more content.")

    jd_keywords = extract_keywords(jd_text, top_n=50)
    result      = compute_ats_score(jd_text, resume_text, jd_keywords)
    result["jd_word_count"]     = len(jd_text.split())
    result["resume_word_count"] = len(resume_text.split())
    return result
